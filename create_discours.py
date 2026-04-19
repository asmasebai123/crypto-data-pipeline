#!/usr/bin/env python3
"""
Génère le document Word avec le discours complet pour la présentation
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Créer le document
doc = Document()

# Configurer les marges
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Ajouter les styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)

# Style Heading 1
heading1_style = doc.styles['Heading 1']
heading1_style.font.name = 'Calibri'
heading1_style.font.size = Pt(16)
heading1_style.font.bold = True
heading1_style.font.color.rgb = RGBColor(6, 90, 130)  # #065A82

# Style Heading 2
heading2_style = doc.styles['Heading 2']
heading2_style.font.name = 'Calibri'
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True
heading2_style.font.color.rgb = RGBColor(28, 114, 147)  # #1C7293

# ===== TITLE =====
title = doc.add_paragraph()
title_run = title.add_run("DISCOURS - CRYPTO DATA PIPELINE")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(6, 90, 130)
title.style = 'Heading 1'

subtitle = doc.add_paragraph("Présentation du Projet Data Engineering End-to-End")
subtitle.paragraph_format.space_after = Pt(18)

# ===== INTRODUCTION =====
doc.add_heading("INTRODUCTION (30 secondes)", level=2)

doc.add_paragraph("Madame, Monsieur le Professeur, chers camarades,")
doc.add_paragraph("Je vous présente aujourd'hui un projet ambitieux : un pipeline Data Engineering complet pour l'analyse des données de cryptomonnaies. C'est un projet qui démontre la maîtrise end-to-end de la chaîne de valeur des données, de l'ingestion jusqu'à la visualisation en production.")
doc.add_paragraph("Tout le pipeline fonctionne 24/7 en cloud et est accessible en ligne. Nous allons voir ensemble comment cela a été réalisé.")

# ===== SLIDE 2: PROBLÉMATIQUE =====
doc.add_heading("SLIDE 2: CONTEXTE ET PROBLÉMATIQUE (1 minute)", level=2)

doc.add_paragraph("Commençons par la problématique qui a motivé ce projet.")
doc.add_paragraph("Le marché des cryptomonnaies est un marché fascinant mais extrêmement volatil. Des milliards de dollars s'échangent chaque jour, 24 heures sur 24. Les prix fluctuent constamment, et les investisseurs ont besoin d'outils d'analyse en temps réel pour prendre des décisions éclairées.")
doc.add_paragraph("Mais le défi ne s'arrête pas là. Nous devons intégrer des données provenant de sources très différentes :")
doc.add_paragraph("• Des APIs REST publiques pour les données batch (CoinGecko)", style='List Bullet')
doc.add_paragraph("• Des flux temps réel via Kafka pour la transmission de masse", style='List Bullet')

doc.add_paragraph("Une fois collectées, ces données doivent être transformées, stockées dans une base relationnelle, puis rendues accessibles à travers un dashboard interactif et professionnel.")
doc.add_paragraph("C'est exactement ce que nous avons construit. Et pour y parvenir, nous avons défini quatre objectifs clairs :")

doc.add_paragraph("1. Un pipeline ETL complet et automatisé", style='List Number')
doc.add_paragraph("2. Une architecture dual-engine capable de traiter à la fois de petits et grands volumes de données", style='List Number')
doc.add_paragraph("3. Un dashboard production-ready avec des visualisations professionnelles", style='List Number')
doc.add_paragraph("4. Un déploiement cloud 24/7 sans nécessiter de maintenance", style='List Number')

# ===== SLIDE 3: ARCHITECTURE =====
doc.add_heading("SLIDE 3: ARCHITECTURE GLOBALE (1 minute 30)", level=2)

doc.add_paragraph("Voici maintenant l'architecture globale de notre système. Elle est organisée en 6 couches bien définies, et c'est justement ce qui en fait sa force.")
doc.add_paragraph("Commençons par le bas :")
doc.add_paragraph("LA COUCHE API : Tout commence avec l'API REST publique de CoinGecko. C'est notre source principale de données pour les informations sur les cryptomonnaies. Nous utilisons aussi Apache Kafka pour simuler des flux temps réel de haute vélocité.")
doc.add_paragraph("LA COUCHE INGESTION : Ici, nous avons deux chemins : un chemin batch qui appelle l'API CoinGecko toutes les 10 minutes, et un chemin streaming qui reçoit les données de Kafka en temps réel.")
doc.add_paragraph("LA COUCHE STOCKAGE : Les données vont dans Neon PostgreSQL, une base de données cloud. C'est ici que tout est persisté. Nous avons plusieurs tables pour les données brutes, les données transformées, et les analyses.")
doc.add_paragraph("LA COUCHE TRANSFORMATION : C'est le cœur du projet. Nous avons implémenté une architecture dual-engine : Pandas pour les petits volumes (développement rapide), et PySpark pour les gros volumes (scalabilité production). Le système choisit automatiquement le meilleur moteur selon le volume.")
doc.add_paragraph("LA COUCHE ORCHESTRATION : Prefect orchestre tout le pipeline, et Railway fait tourner le code 24/7 en cloud sans qu'on ait besoin de faire quoi que ce soit.")
doc.add_paragraph("LA COUCHE VISUALISATION : Enfin, Streamlit Cloud expose un dashboard interactif accessible publiquement. C'est ce que vos utilisateurs finaux verront.")

# ===== SLIDE 4: INGESTION BATCH =====
doc.add_heading("SLIDE 4: INGESTION BATCH (45 secondes)", level=2)

doc.add_paragraph("Parlons d'abord de l'ingestion batch. Toutes les 10 minutes, notre système appelle l'API CoinGecko et demande les données actuelles pour 5 cryptomonnaies : Bitcoin, Ethereum, Solana, XRP et BNB.")
doc.add_paragraph("Chaque réponse nous donne le prix actuel, la capitalisation boursière, le volume d'échange 24 heures, et les variations en pourcentage.")
doc.add_paragraph("CoinGecko a été choisi parce que c'est l'une des sources les plus fiables du marché, avec une API gratuite bien documentée. La fréquence de 10 minutes respecte les limites du plan gratuit (30 appels/minute) tout en maintenant une granularité suffisante pour l'analyse.")
doc.add_paragraph("Vous voyez sur la droite un exemple de réponse JSON. C'est simple et structuré.")

# ===== SLIDE 5: INGESTION STREAMING =====
doc.add_heading("SLIDE 5: INGESTION STREAMING KAFKA (45 secondes)", level=2)

doc.add_paragraph("Maintenant, Kafka. C'est un système de messaging distribué, le standard industrie pour gérer des flux de données en temps réel.")
doc.add_paragraph("Dans notre système, nous avons un topic Kafka appelé 'crypto_prices'. Un producteur envoie des messages toutes les 5 secondes avec le prix actuel et la variation de chaque crypto. Un consommateur reçoit ces messages et les stocke.")
doc.add_paragraph("Pourquoi Kafka ? Parce que c'est l'architecture qu'on trouve dans la majorité des systèmes data en production. C'est hautement scalable et peut gérer des millions de messages par seconde. Dans notre cas, nous simulons les données, mais en production, on pourrait connecter directement aux exchanges cryptographiques via WebSocket.")
doc.add_paragraph("Cela démontre l'architecture production-ready que nous avons construite.")

# ===== SLIDE 6: STOCKAGE =====
doc.add_heading("SLIDE 6: STOCKAGE POSTGRESQL NEON (1 minute)", level=2)

doc.add_paragraph("Toutes les données, qu'elles viennent de batch ou streaming, vont dans Neon PostgreSQL Cloud.")
doc.add_paragraph("Nous avons plusieurs tables :")

doc.add_paragraph("'raw_crypto_prices' : les données brutes de l'API batch", style='List Bullet')
doc.add_paragraph("'stream_crypto_prices' : les données du flux Kafka", style='List Bullet')
doc.add_paragraph("Des tables transformées pour les moyennes horaires, les classements, les alertes", style='List Bullet')
doc.add_paragraph("Des tables spécialisées pour les analyses Spark SQL avancées", style='List Bullet')

doc.add_paragraph("Pourquoi Neon ? C'est un PostgreSQL serverless gratuit, hautement disponible, avec sauvegarde automatique. Nous n'avons pas à gérer de serveur. C'est la combinaison parfaite pour un projet de démonstration, mais aussi pour la production.")
doc.add_paragraph("Les données persistent ici de façon sécurisée et accessible.")

# ===== SLIDE 7: TRANSFORMATION =====
doc.add_heading("SLIDE 7: TRANSFORMATION DUAL-ENGINE (2 minutes)", level=2)

doc.add_paragraph("Voici l'innovation clé de ce projet : l'architecture dual-engine.")
doc.add_paragraph("Le problème classique en data engineering : comment traiter efficacement TOUS les volumes ? Les petits volumes ? Les grands volumes ?")
doc.add_paragraph("Notre réponse : deux moteurs, deux approches.")
doc.add_paragraph("PANDAS : quand les données sont en nombre limité (moins de 50 000 lignes), Pandas est plus rapide, plus simple, et plus agréable à développer. C'est du Python pur, facile à déboguer.")
doc.add_paragraph("PYSPARK : quand les données explosent (50 000 lignes ou plus), nous basculons sur Spark. Spark peut traiter des milliards de lignes distribuées sur un cluster. C'est production-ready.")
doc.add_paragraph("Et la beauté ? Le système choisit AUTOMATIQUEMENT le bon moteur selon le volume de données. Pas de changement de code.")
doc.add_paragraph("Nous avons implémenté 5 transformations :")

doc.add_paragraph("1. NETTOYAGE : suppression des valeurs nulles, doublons, et anomalies", style='List Number')
doc.add_paragraph("2. MOYENNE HORAIRE : agrégation du prix moyen, min et max par heure", style='List Number')
doc.add_paragraph("3. CLASSEMENT JOURNALIER : rangement des cryptos par performance 24h. Là, Spark brille vraiment avec ses Window Functions.", style='List Number')
doc.add_paragraph("4. DÉTECTION D'ALERTES : quand la variation dépasse ±5%, on génère une alerte", style='List Number')
doc.add_paragraph("5. VOLUME JOURNALIER : somme des volumes échangés par jour", style='List Number')

# ===== SLIDE 8: ORCHESTRATION =====
doc.add_heading("SLIDE 8: ORCHESTRATION PREFECT + RAILWAY (1 minute 30)", level=2)

doc.add_paragraph("Comment s'assurer que tout cela tourne correctement, sans intervention humaine, 24 heures par jour ?")
doc.add_paragraph("C'est le rôle de Prefect et Railway.")
doc.add_paragraph("PREFECT : c'est un framework d'orchestration moderne. Nous définissons un 'flow' qui exécute les étapes dans l'ordre :")

doc.add_paragraph("1. Créer les tables PostgreSQL", style='List Number')
doc.add_paragraph("2. Appeler l'API CoinGecko", style='List Number')
doc.add_paragraph("3. Sauvegarder les données brutes", style='List Number')
doc.add_paragraph("4. Exécuter les transformations (Pandas ou Spark)", style='List Number')
doc.add_paragraph("5. Générer les analyses avancées", style='List Number')
doc.add_paragraph("6. Rafraîchir le dashboard", style='List Number')

doc.add_paragraph("Prefect gère les dépendances, les retries automatiques en cas d'erreur, et l'logging.")
doc.add_paragraph("RAILWAY : c'est la plateforme cloud qui fait tourner ce pipeline 24/7. Railway lance le code, le relance automatiquement en cas de crash, et tout cela coûte pratiquement rien : 5 dollars de crédit par mois, suffisant pour nos besoins.")
doc.add_paragraph("Résultat : un pipeline qui s'exécute toutes les 10 minutes, jour et nuit, sans que nous ayons à faire quoi que ce soit.")

# ===== SLIDE 9: DASHBOARD =====
doc.add_heading("SLIDE 9: DASHBOARD STREAMLIT (1 minute 30)", level=2)

doc.add_paragraph("Alors, où finissent les données ? Dans un dashboard. Et pas n'importe quel dashboard.")
doc.add_paragraph("Nous avons construit une interface avec 5 onglets différents, chacun offrant une perspective unique sur les données :")

doc.add_paragraph("ONGLET 1 - VUE D'ENSEMBLE : les prix actuels en cartes colorées, un graphique montrant l'évolution, et le classement des meilleures performances du jour", style='List Bullet')
doc.add_paragraph("ONGLET 2 - ANALYSE TECHNIQUE : pour les traders. Candlestick OHLC, heatmap de corrélation entre les cryptos", style='List Bullet')
doc.add_paragraph("ONGLET 3 - ALERTES & VOLUME : affiche les cryptos qui ont des variations importantes, et les volumes d'échange", style='List Bullet')
doc.add_paragraph("ONGLET 4 - STREAMING KAFKA : pour voir les flux temps réel, les prix s'actualisent toutes les 5 secondes", style='List Bullet')
doc.add_paragraph("ONGLET 5 - SPARK ANALYTICS : analyses avancées calculées via SQL : la volatilité de chaque crypto, et sa part dans le marché global", style='List Bullet')

doc.add_paragraph("Le dashboard a un design moderne avec un dark theme professionnel. C'est accessible 24/7 sur internet. Et il se rafraîchit automatiquement toutes les 60 secondes.")
doc.add_paragraph("C'est du niveau production.")

# ===== SLIDE 10: RÉSULTATS =====
doc.add_heading("SLIDE 10: RÉSULTATS ATTENDUS VS RÉALITÉ (1 minute)", level=2)

doc.add_paragraph("Parlons de ce que nous avions prévu, et de ce que nous avons réellement livré.")

p = doc.add_paragraph()
p.add_run("ATTENDUS :").bold = True

doc.add_paragraph("Un pipeline ETL complet", style='List Bullet')
doc.add_paragraph("Environ 5 000 lignes de données par jour", style='List Bullet')
doc.add_paragraph("Une architecture dual-engine Pandas/Spark", style='List Bullet')
doc.add_paragraph("5 transformations métier", style='List Bullet')
doc.add_paragraph("Un dashboard interactif", style='List Bullet')
doc.add_paragraph("Orchestration et exécution automatiques", style='List Bullet')
doc.add_paragraph("Déploiement cloud", style='List Bullet')

p = doc.add_paragraph()
p.add_run("RÉALITÉ :").bold = True

doc.add_paragraph("Pipeline complet et fonctionnel ✓", style='List Bullet')
doc.add_paragraph("5 050 lignes effectivement stockées ✓", style='List Bullet')
doc.add_paragraph("Architecture dual-engine opérationnelle ✓", style='List Bullet')
doc.add_paragraph("5 transformations + 3 analyses Spark SQL avancées ✓", style='List Bullet')
doc.add_paragraph("Dashboard version 2.1 avec 5 onglets ✓", style='List Bullet')
doc.add_paragraph("Prefect + Railway fonctionnant 24/7 ✓", style='List Bullet')
doc.add_paragraph("Neon + Streamlit Cloud en production ✓", style='List Bullet')

# ===== SLIDE 11: MÉTRIQUES =====
doc.add_heading("SLIDE 11: DONNÉES DE SORTIE (45 secondes)", level=2)

doc.add_paragraph("Voici les chiffres clés du projet :")

doc.add_paragraph("5 cryptomonnaies suivies", style='List Bullet')
doc.add_paragraph("~5 000 lignes collectées par jour", style='List Bullet')
doc.add_paragraph("Fréquence batch : 10 minutes", style='List Bullet')
doc.add_paragraph("Fréquence streaming : 5 secondes", style='List Bullet')
doc.add_paragraph("8 tables PostgreSQL créées", style='List Bullet')
doc.add_paragraph("5 transformations métier", style='List Bullet')
doc.add_paragraph("3 analyses Spark SQL avancées", style='List Bullet')
doc.add_paragraph("Uptime : 24/7", style='List Bullet')

# ===== SLIDE 12: FONCTIONNALITÉS =====
doc.add_heading("SLIDE 12: FONCTIONNALITÉS DÉLIVRÉES (1 minute)", level=2)

doc.add_paragraph("Récapitulons ce que ce projet offre :")

doc.add_paragraph("PIPELINE AUTOMATISÉ : Une fois lancé, il s'exécute tous les jours sans intervention humaine, 24 heures sur 24.", style='List Number')
doc.add_paragraph("DUAL-ENGINE : Flex automatiquement entre Pandas (efficace) et Spark (scalable) selon le volume.", style='List Number')
doc.add_paragraph("DASHBOARD INTERACTIF : Interface utilisateur moderne avec 5 onglets thématiques, dark theme professionnel.", style='List Number')
doc.add_paragraph("ANALYSES AVANCÉES : Volatilité, dominance marché, corrélation entre assets. Tout calculé via SQL, pas besoin de code complexe.", style='List Number')
doc.add_paragraph("CLOUD DEPLOYMENT : Entièrement hébergé. Railway para l'orchestration, Neon pour la base de données, Streamlit Cloud pour le dashboard. Tout gratuit ou quasi-gratuit.", style='List Number')
doc.add_paragraph("ALERTES TEMPS RÉEL : Détection automatique des mouvements de prix importants.", style='List Number')

# ===== SLIDE 13: TECH STACK =====
doc.add_heading("SLIDE 13: STACK TECHNOLOGIQUE (30 secondes)", level=2)

doc.add_paragraph("Pour ceux intéressés par les détails techniques, voici les technologies utilisées :")

doc.add_paragraph("INGESTION : Python, API REST CoinGecko, Kafka", style='List Bullet')
doc.add_paragraph("STOCKAGE : PostgreSQL 15, Neon Cloud", style='List Bullet')
doc.add_paragraph("TRANSFORMATION : Pandas 2.2.3, PySpark 3.5.1, SQL", style='List Bullet')
doc.add_paragraph("ORCHESTRATION : Prefect 3.x, Railway", style='List Bullet')
doc.add_paragraph("VISUALISATION : Streamlit 1.42, Plotly 5.24, Streamlit Cloud", style='List Bullet')
doc.add_paragraph("INFRASTRUCTURE : Docker, Python 3.12, GitHub", style='List Bullet')

# ===== CONCLUSION =====
doc.add_heading("SLIDE 14: CONCLUSION (1 minute 30)", level=2)

doc.add_paragraph("En conclusion, ce projet démontre une maîtrise complète de la chaîne de valeur Data Engineering.")
doc.add_paragraph("Nous avons ingéré des données depuis plusieurs sources. Nous les avons transformées en utilisant deux moteurs différents pour maximiser l'efficacité. Nous les avons stockées de façon sécurisée et scalable. Et nous les avons exposées par un interface moderne et accessible.")
doc.add_paragraph("Ce qui est remarquable, c'est que le TOUT fonctionne automatiquement en cloud, sans intervention humaine. C'est cela, la production-ready.")
doc.add_paragraph("Les compétences développées incluent :")

doc.add_paragraph("Ingestion batch et streaming", style='List Bullet')
doc.add_paragraph("Stockage relationnel cloud", style='List Bullet')
doc.add_paragraph("Transformations Pandas ET PySpark", style='List Bullet')
doc.add_paragraph("Spark SQL, Window Functions", style='List Bullet')
doc.add_paragraph("Orchestration avec Prefect", style='List Bullet')
doc.add_paragraph("Déploiement cloud (Railway, Neon, Streamlit)", style='List Bullet')
doc.add_paragraph("Visualisation interactive professionnel", style='List Bullet')

doc.add_paragraph("Le dashboard est actuellement en ligne et fonctionne parfaitement. Si vous êtes intéressés, vous pouvez le visiter pour voir les données en direct.")
doc.add_paragraph("Merci de votre attention. Je suis prête pour vos questions.")

# Sauvegarder le document
doc.save('/sessions/peaceful-sweet-wright/mnt/crypto_data_project/discours_presentation.docx')
print("✅ Discours créé avec succès: discours_presentation.docx")
